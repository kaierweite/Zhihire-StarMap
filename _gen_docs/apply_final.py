import os, json
from docx import Document

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
m = json.load(open(os.path.join("_gen_docs", "prod_final.json"), "r", encoding="utf-8"))

doc = Document(os.path.join("docs", "软件产品说明书.docx"))

count = 0
for para in doc.paragraphs:
    for old, new in m.items():
        if para.text.strip() == old.strip():
            para.clear()
            para.add_run(new)
            count += 1
            break

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in m.items():
                    if para.text.strip() == old.strip():
                        para.clear()
                        para.add_run(new)
                        count += 1
                        break

doc.save(os.path.join("docs", "软件产品说明书.docx"))
print(f"Product manual: {count} final replacements applied!")