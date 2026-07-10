# -*- coding: utf-8 -*-
import os
from docx import Document

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")

def update_basic(doc_path):
    doc = Document(os.path.join("docs", doc_path))
    updates = {
        "XX（软件名称）": "智聘星图 (Zhihire StarMap)",
        "2026-07-01": "2026-07-10",
        "张三": "开发组",
        "李四": "测试组",
        "王五": "开发组",
    }
    for para in doc.paragraphs:
        for old, new in updates.items():
            if para.text.strip() == old.strip():
                para.clear()
                para.add_run(new)
                break
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in updates.items():
                        if para.text.strip() == old.strip():
                            para.clear()
                            para.add_run(new)
                            break
    doc.save(os.path.join("docs", doc_path))
    print(f"Updated: {doc_path}")

update_basic("软件功能测试报告.docx")
update_basic("软件功能设计文档.docx")
update_basic("软件产品说明书.docx")