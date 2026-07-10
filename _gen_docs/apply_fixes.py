import os, json
from docx import Document

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
base = os.path.join("docs")

fix_map = json.load(open(os.path.join("_gen_docs", "fix_map.json"), "r", encoding="utf-8"))

for doc_name in ["软件功能测试报告.docx", "软件功能设计文档.docx", "软件产品说明书.docx"]:
    doc = Document(os.path.join(base, doc_name))
    count = 0
    for para in doc.paragraphs:
        for old, new in fix_map.items():
            if para.text.strip() == old.strip():
                para.clear()
                para.add_run(new)
                count += 1
                break
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in fix_map.items():
                        if para.text.strip() == old.strip():
                            para.clear()
                            para.add_run(new)
                            count += 1
                            break
    doc.save(os.path.join(base, doc_name))
    print(f"{doc_name}: {count} fixes applied")