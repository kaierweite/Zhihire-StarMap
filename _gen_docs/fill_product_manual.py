# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
doc = Document(os.path.join("docs", "软件产品说明书.docx"))

replacements = {
    "软件产品说明书（用户手册）": "软件产品说明书（用户手册）",
    'XX（软件名称）': "智聘星图 (Zhihire StarMap)",
    "【填写示例】 本文档为《校园图书借阅管理系统》的官方用户使用说明，面向全体终端用户与系统管理员，旨在帮助用户快速了解产品功能、掌握操作方法、解决常见使用问题，同时作为产品交付与验收的配套说明文件。": "本文档为《智聘星图 (Zhihire StarMap)》的官方用户使用说明，面向求职者、企业用户与系统管理员三类角色，旨在帮助用户快速了解产品功能、掌握操作方法、解决常见使用问题，同时作为产品交付与验收的配套说明文件。\n\n本系统是第十五届中国软件杯大赛 B2 赛题作品，基于银河麒麟操作系统与 KingbaseES 数据库构建，集成 DeepSeek 大模型 AI 能力，实现简历智能解析、人岗精准匹配、能力图谱展示与 AI 模拟面试等核心功能。",
    "【填写示例】\n适用用户：在校师生、图书馆管理员、系统运维人员；\n适用版本：系统 V1.0 正式版本；\n适用场景：日常图书借阅业务、图书馆运营管理、系统基础运维。": "适用用户：求职者（个人用户）、企业HR/招聘负责人、系统管理员；\n适用版本：系统 V1.0 正式版本；\n适用场景：求职者简历管理与 AI 解析、岗位搜索与智能匹配、能力图谱与职业规划、AI 模拟面试训练、企业岗位发布与候选人筛选、管理员系统运维与审核管理。",
    "GB/T 8567-2006 计算机软件文档编制规范\nGB/T 25000.51-2016 就绪可用软件产品质量要求和测试细则\n《校园图书借阅 Web 系统功能需求分析文档 V1.0》\n产品正式版功能清单": "GB/T 8567-2006 计算机软件文档编制规范\nGB/T 25000.51-2016 就绪可用软件产品质量要求和测试细则\n《智聘星图功能需求分析文档 V1.0》\n智聘星图产品正式版功能清单\n第十五届中国软件杯大赛 B2 赛题说明书",
    "馆藏": "技能图谱",
    "图书馆收藏的全部图书总量": "以类型化边（PREREQUISITE/RELATED/EXTENDS）构建的技能知识图谱，反映技能间的依赖关系",
    "ISBN": "JD",
    "图书的唯一编号，类似图书的「身份证号」": "Job Description，岗位需求描述文档",
    "V1.0": "V1.0",
    "2026-07-01": "2026-07-10",
    "对应系统 V1.0 版本，发布首版产品说明书": "对应系统 V1.0 版本，发布首版产品说明书",
    "张三": "开发组",
    "《校园图书借阅 Web 系统功能需求分析文档 V1.0》": "《智聘星图功能需求分析文档 V1.0》\n智聘星图产品正式版功能清单\n第十五届中国软件杯大赛 B2 赛题说明书",
    "【填写示例】 智聘星图（以下简称"本系统"）是基于银河麒麟操作系统的 AI 智能匹配与能力图谱平台。系统面向求职者与招聘企业，提供简历智能解析、个人能力图谱生成、人岗智能匹配推荐、AI 职业规划及 AI 模拟面试等核心功能，解决传统招聘流程中简历筛选效率低、人岗匹配不精准、求职者职业发展路径不清晰等痛点。": "智聘星图（以下简称"本系统"）是基于银河麒麟操作系统的 AI 智能匹配与能力图谱平台。系统面向求职者与招聘企业，提供简历智能解析、个人能力图谱生成、人岗智能匹配推荐、AI 职业规划及 AI 模拟面试等核心功能，解决传统招聘流程中简历筛选效率低、人岗匹配不精准、求职者职业发展路径不清晰等痛点。",
    "【填写示例】\n系统覆盖求职者端、企业端、管理端三大模块：\n求职者端：图书检索、图书借阅预约、个人借阅记录查询、个人信息管理\n企业端：图书信息管理、借还登记操作、读者信息管理\n管理端：系统配置维护、数据统计查看、操作日志查询": "系统覆盖求职者端、企业端、管理端三大模块：\n\n求职者端：\n1. 用户注册登录 - 支持求职者与企业双角色注册，JWT 认证\n2. 个人档案管理 - 编辑个人资料、技能标签、教育/工作经历\n3. 简历中心 - 简历上传（PDF/DOC/DOCX）、AI 智能解析、简历优化\n4. 能力图谱 - 个人技能知识图谱（ECharts 力导向图），含缺口视图\n5. 岗位推荐 - AI 智能匹配推荐，多维度评分与推荐依据\n6. 岗位搜索 - 按关键词/薪资/学历/地点多维筛选\n7. 职业规划 - AI 生成学习路径与技能补全建议\n8. AI 模拟面试 - LLM 驱动的面试官，即时评分与面试报告\n9. 通知中心 - 站内消息通知\n\n企业端：\n1. 企业注册与信息管理\n2. 岗位发布 - 手动填写/JD 上传双模式，AI 解析自动填充\n3. 岗位管理 - 列表查看、编辑、上下架\n4. 候选人推荐 - 反向匹配推荐，匹配度排序\n5. 智能筛选 - 人才搜索与对比\n6. 岗位能力图谱 - 查看岗位技能要求图谱\n7. 通知中心\n\n管理端：\n1. 仪表板 - 实时数据统计（用户/岗位/企业汇总）\n2. 用户管理 - 搜索、查看、封禁/解封\n3. 审核管理 - 企业资质审核、技能字典审核\n4. 系统日志 - 操作日志查询\n5. AI 模型配置 - DeepSeek API 参数配置",
    "【填写示例】\n- AI 智能简历解析：纯 Python 实现的 PDF/DOCX 解析引擎，适配 LoongArch 龙芯架构，无需外部 API\n- 类型化边知识图谱：networkx 内存图谱，PREREQUISITE/RELATED/EXTENDS 三种关系，按 skill.category 上色\n- 可解释匹配评分：维度子分 + 匹配依据自然语言生成，求职者知道"为什么匹配"和"差在哪里"\n- AI 职业规划：图算法缺口分析 + LLM 润色，生成有序学习路径与资源推荐\n- AI 模拟面试：LLM 驱动面试官，支持文字/语音多模态交互，即时评分与报告\n- 国产化全栈：银河麒麟 V11 + KingbaseES + LoongArch 龙芯，全链路国产适配": "AI 智能简历解析：纯 Python 实现的 PDF/DOCX 解析引擎（pdfplumber + python-docx），适配 LoongArch 龙芯架构\n类型化边知识图谱：networkx 常驻内存图谱，PREREQUISITE/RELATED/EXTENDS 三种关系，按 skill.category 分色展示\n可解释匹配评分：维度子分 + 匹配依据自然语言生成，求职者知道"为什么匹配"和"差在哪里"\nAI 职业规划：图算法缺口分析 + LLM 润色，生成有序学习路径与学习资源推荐\nAI 模拟面试：DeepSeek LLM 驱动面试官，即时评分与多维度面试报告\n国产化全栈：银河麒麟 V11 + KingbaseES 数据库 + LoongArch 龙芯，全链路国产适配",
    "【填写示例】\n系统采用 HTTPS 协议传输数据，密码使用 BCrypt 加密存储，接口通过 JWT Token 鉴权。\n会话超时时间：30 分钟\n数据传输：全程 HTTPS 加密\n密码策略：长度 8-64 位，字母+数字组合\n推荐并发用户数：50-100 人同时在线": "系统采用 BCrypt 密码加密存储，前端与后端接口通过 JWT Token 鉴权，支持跨域 CORS 配置。\n\n会话超时时间：24 小时（JWT Token 有效期）\n数据加密：BCrypt 密码加密、JWT Token 加密传输\n密码策略：长度 8-64 位，字母+数字组合\n推荐并发用户数：50-100 人同时在线\n\n本系统部署于银河麒麟 V11 操作系统，基于 LoongArch（龙芯）架构，使用 KingbaseES 数据库，确保数据安全可控。",
    "【填写示例】\n服务器：4 核 8GB 内存，100GB 硬盘，银河麒麟 V11 操作系统\n客户端：Windows 10/11 或 macOS，Chrome/Firefox/Edge 最新版本\n网络：100Mbps 以上局域网或互联网连接": "服务器端：\n- 操作系统：银河麒麟高级服务器版 V11\n- CPU 架构：LoongArch（龙芯），4 核及以上\n- 内存：8GB 及以上\n- 硬盘：256GB 及以上\n- 数据库：人大金仓 KingbaseES V8（兼容 PostgreSQL 模式）\n\n客户端：\n- 操作系统：Windows 10/11、macOS、Linux 桌面版\n- 浏览器：Chrome/Firefox/Edge 最新正式版本\n- 屏幕分辨率：1280x720 及以上推荐",
    "【填写示例】\n当前版本暂无移动端专属适配，建议在平板或大屏手机上使用浏览器访问系统。": "当前版本暂无移动端专属适配，但基于 Element Plus 响应式布局，可在平板设备上正常使用核心功能。建议在 1280x768 及以上分辨率的桌面端浏览器获得最佳体验。",
    "【填写示例】\n推荐带宽：10Mbps 以上\n需要访问的域名/IP：后端 API 服务地址（由管理员提供）\n外网访问：需要配置公网 IP 或域名解析": "推荐带宽：10Mbps 以上\n需要访问的域名/IP：后端 API 服务地址（由管理员提供，默认 http://localhost:8000）\n外网访问：需要配置公网 IP 或域名解析\nAI 服务依赖：需要访问 DeepSeek API（https://api.deepseek.com）",
    "【填写示例】\n浏览器地址栏输入 http://localhost:8080 或管理员提供的访问地址\n系统自动跳转至首页，无需手动安装任何插件或客户端": "浏览器地址栏输入管理员提供的访问地址（开发环境默认 http://localhost:5173）\n系统自动跳转至首页，无需手动安装任何插件或客户端\n\n管理员可访问 Swagger 接口文档：http://localhost:8000/api/docs\nReDoc 文档：http://localhost:8000/api/redoc",
    "【填写示例】\n管理员账号：admin\n初始密码：admin123\n首次登录后请立即修改密码": "管理员账号在系统初始化时通过数据库脚本创建\n\n默认管理员信息：\n- 用户名：admin\n- 初始密码：由部署脚本生成，首次登录后请立即修改\n\n求职者与企业用户可通过注册页自行注册。",
    "【填写示例】\n管理员登录后建议依次完成以下初始配置：\n1. 进入【系统配置】页面配置图书馆借阅规则（借阅天数、续借次数等）\n2. 在【读者管理】中添加或导入读者数据\n3. 在【图书管理】中添加图书信息或批量导入\n4. 检查系统运行状态是否正常": "管理员登录后建议依次完成以下初始配置：\n1. 进入【AI 模型配置】页面配置 DeepSeek API Key\n2. 在【审核管理】中审核待处理的企业注册申请\n3. 在【用户管理】中检查系统用户列表\n4. 检查系统仪表板数据统计是否正常",
}

def replace_text_in_paragraph(para, replacements_dict):
    if para.text in replacements_dict:
        para.clear()
        run = para.add_run(replacements_dict[para.text])
        run.font.size = para.runs[0].font.size if para.runs else Pt(10.5)
        run.font.name = para.runs[0].font.name if para.runs else "宋体"
        return True
    return False

for para in doc.paragraphs:
    replace_text_in_paragraph(para, replacements)

# Update tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text == 'XX（软件名称）':
                    para.clear()
                    run = para.add_run("智聘星图 (Zhihire StarMap)")
                elif '校园图书借阅' in para.text and '需求' in para.text:
                    para.clear()
                    run = para.add_run("《智聘星图功能需求分析文档 V1.0》")

# Fix version table
for table in doc.tables:
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                if para.text == '2026-07-01':
                    para.clear()
                    run = para.add_run("2026-07-10")
                if para.text == '张三' and len(row.cells) > 3:
                    para.clear()
                    run = para.add_run("开发组")

# Save
outpath = os.path.join("docs", "软件产品说明书.docx")
doc.save(outpath)
print(f"Product manual saved to {outpath}")
