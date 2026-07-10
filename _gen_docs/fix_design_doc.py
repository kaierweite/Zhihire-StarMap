import os, json
from docx import Document

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
m = json.load(open(os.path.join("_gen_docs", "design_fix.json"), "r", encoding="utf-8"))

doc = Document(os.path.join("docs", "软件功能设计文档.docx"))

# Apply paragraph fixes
pcount = 0
for para in doc.paragraphs:
    for old, new in m.items():
        if old and para.text.strip() == old.strip():
            para.clear()
            if new:
                para.add_run(new)
            pcount += 1
            break

# Fix Table 1 - Terms: remove MyBatis, add correct terms
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip() == "MyBatis":
                row.cells[0].text = "SQLAlchemy 2.0"
                row.cells[1].text = "Python异步ORM框架，用于KingbaseES数据库的声明式映射与查询"
                pcount += 1
            if "Java 持久层框架" in cell.text:
                row.cells[1].text = "Python异步ORM框架，用于KingbaseES数据库的声明式映射与查询"
                pcount += 1

# Fix Table 3 - Parameter table: replace book search params with resume parse
for table in doc.tables:
    if len(table.rows) == 4 and table.rows[0].cells[0].text == "参数名":
        if table.rows[1].cells[0].text == "keyword":
            table.rows[1].cells[0].text = "file"
            table.rows[1].cells[1].text = "UploadFile"
            table.rows[1].cells[2].text = "是"
            table.rows[1].cells[3].text = "待解析的简历文件，仅.pdf/.doc/.docx"
            table.rows[2].cells[0].text = "resume_id"
            table.rows[2].cells[1].text = "Integer"
            table.rows[2].cells[2].text = "是"
            table.rows[2].cells[3].text = "上传后返回的简历ID，用于查询解析结果"
            table.rows[3].cells[0].text = "task_id"
            table.rows[3].cells[1].text = "Integer"
            table.rows[3].cells[2].text = "否"
            table.rows[3].cells[3].text = "异步解析任务ID(不传则同步解析)"
            pcount += 1

# Fix Table 4 - Return params: book list -> resume parse result
for table in doc.tables:
    if len(table.rows) == 5 and table.rows[0].cells[0].text == "参数名":
        if table.rows[1].cells[0].text == "code":
            table.rows[1].cells[0].text = "code"
            table.rows[1].cells[1].text = "Integer"
            table.rows[1].cells[2].text = "响应状态码，200成功"
            table.rows[2].cells[0].text = "message"
            table.rows[2].cells[1].text = "String"
            table.rows[2].cells[2].text = "响应提示信息"
            table.rows[3].cells[0].text = "data.name"
            table.rows[3].cells[1].text = "String"
            table.rows[3].cells[2].text = "解析出的姓名"
            table.rows[4].cells[0].text = "data.skills"
            table.rows[4].cells[1].text = "Array<String>"
            table.rows[4].cells[2].text = "解析出的技能列表(已归一化)"
            pcount += 1

# Fix Table 5 - Database table: book -> user table
for table in doc.tables:
    if len(table.rows) == 11 and table.rows[0].cells[0].text == "字段名":
        if table.rows[1].cells[0].text == "book_id":
            rows_data = [
                ["user_id", "bigint", "主键、自增", "用户唯一ID"],
                ["username", "varchar(32)", "非空、唯一索引", "用户名，4-32字符"],
                ["password_hash", "varchar(128)", "非空", "BCrypt加密密码"],
                ["role", "varchar(20)", "非空、索引", "角色: USER/COMPANY/ADMIN"],
                ["status", "varchar(20)", "非空、默认NORMAL", "状态: NORMAL/DISABLED/BANNED"],
                ["email", "varchar(100)", "可空", "电子邮箱"],
                ["phone", "varchar(20)", "可空", "手机号码"],
                ["nickname", "varchar(50)", "可空", "用户昵称"],
                ["created_at", "datetime", "非空", "创建时间"],
                ["updated_at", "datetime", "非空", "更新时间"],
            ]
            for ri, row_data in enumerate(rows_data, 1):
                for ci, val in enumerate(row_data):
                    table.rows[ri].cells[ci].text = val
            pcount += 1

# Fix Table 6 - Version: update V1.1
for table in doc.tables:
    if len(table.rows) == 2 and table.rows[0].cells[0].text == "变更版本":
        # Already has AI interview module update, just update date
        table.rows[1].cells[1].text = "2026-07-10"
        pcount += 1

doc.save(os.path.join("docs", "软件功能设计文档.docx"))
print(f"Design doc: {pcount} fixes applied!")