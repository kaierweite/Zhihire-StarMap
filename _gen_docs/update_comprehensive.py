import os
from docx import Document
from docx.shared import Inches

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
img_path = r"docs\images\architecture.png"

def replace_text(doc, updates):
    for para in doc.paragraphs:
        for old, new in updates.items():
            if para.text.strip() == old.strip():
                para.clear()
                para.add_run(new)
                break
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in updates.items():
                        if para.text.strip() == old.strip():
                            para.clear()
                            para.add_run(new)
                            break

def update_product_manual():
    doc = Document(os.path.join("docs", "软件产品说明书.docx"))
    replace_text(doc, {
        "图书馆收藏的全部图书总量": "以类型化边(PREREQUISITE/RELATED/EXTENDS)构建的技能知识图谱",
        "图书的唯一编号，类似图书的「身份证号」": "Job Description/岗位需求描述文档的缩写",
        "单次借书可使用的最长天数": "AI面试会话从创建到完成的时间限制",
        "借书到期前，延长借阅时长的操作": "在AI面试中请求换题或重答的操作",
    })
    doc.save(os.path.join("docs", "软件产品说明书.docx"))
    print("Product manual: terms updated")

def update_test_report():
    doc = Document(os.path.join("docs", "软件功能测试报告.docx"))
    replace_text(doc, {
        "《校园图书借阅 Web 系统功能需求分析文档 V1.0》": "《智聘星图功能需求分析文档 V1.0》",
        "校园图书借阅 Web 系统 V1.0 正式版本": "智聘星图 (Zhihire StarMap) V1.0 正式版本",
        "读者端模块": "求职者端模块",
        "管理员端模块": "企业管理端模块",
        "系统管理模块": "管理后台模块",
        "4 核 8G 云服务器，Windows Server 2019，JDK 1.8，": "4核8G内存，银河麒麟V11，Python 3.12，KingbaseES V8",
        "Windows 11 操作系统，16G 内存，Chrome 120、Edge 1": "Windows 11/银河麒麟桌面版，Chrome/Firefox/Edge最新版",
    })
    doc.save(os.path.join("docs", "软件功能测试报告.docx"))
    print("Test report: basic info updated")

def update_design_doc():
    doc = Document(os.path.join("docs", "软件功能设计文档.docx"))
    replace_text(doc, {
        "《校园图书借阅 Web 系统功能需求分析文档 V1.0》": "《智聘星图功能需求分析文档 V1.0》",
        "SpringBoot 2.x 开发规范手册": "FastAPI + Python 3.12 开发规范",
        "MySQL 数据库设计规范": "KingbaseES 数据库设计规范",
        "B/S 架构": "B/S(浏览器/服务器)架构",
        "浏览器 / 服务器架构，用户通过浏览器即可访问系统，无需安装客户端": "浏览器/服务器架构，用户通过浏览器访问系统，Vue3前端+FastAPI后端",
        "基于角色的权限访问控制模型，通过角色关联用户与权限": "基于角色(User/Company/Admin)的JWT权限控制模型",
    })
    doc.save(os.path.join("docs", "软件功能设计文档.docx"))
    print("Design document: basic info updated")

update_product_manual()
update_test_report()
update_design_doc()