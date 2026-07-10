# -*- coding: utf-8 -*-
import os
from docx import Document

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
doc = Document(os.path.join("docs", "软件产品说明书.docx"))

updates = {
    "XX（软件名称）": "智聘星图 (Zhihire StarMap)",
    "2026-07-01": "2026-07-10",
    "张三": "开发组",
}

for para in doc.paragraphs:
    for old, new in updates.items():
        if para.text.strip() == old.strip():
            para.clear()
            para.add_run(new)
            break

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in updates.items():
                    if para.text.strip() == old.strip():
                        para.clear()
                        para.add_run(new)
                        break

outpath = os.path.join("docs", "软件产品说明书.docx")
doc.save(outpath)
print("Product manual basic update OK!")