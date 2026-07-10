# -*- coding: utf-8 -*-
import os, json
from docx import Document

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
base = os.path.join("docs")

def apply_replacements(doc_path, mapping):
    doc = Document(os.path.join(base, doc_path))
    count = 0
    for para in doc.paragraphs:
        for old, new in mapping.items():
            if para.text.strip() == old.strip():
                para.clear()
                para.add_run(new)
                count += 1
                break
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in mapping.items():
                        if para.text.strip() == old.strip():
                            para.clear()
                            para.add_run(new)
                            count += 1
                            break
    doc.save(os.path.join(base, doc_path))
    return count

# Product manual
pm = json.load(open(os.path.join("_gen_docs", "prod_map.json"), "r", encoding="utf-8"))
c1 = apply_replacements("软件产品说明书.docx", pm)
print(f"Product manual: {c1} replacements")

# Test report
tm = json.load(open(os.path.join("_gen_docs", "test_map.json"), "r", encoding="utf-8"))
c2 = apply_replacements("软件功能测试报告.docx", tm)
print(f"Test report: {c2} replacements")

# Design doc
dm = json.load(open(os.path.join("_gen_docs", "design_map.json"), "r", encoding="utf-8"))
c3 = apply_replacements("软件功能设计文档.docx", dm)
print(f"Design doc: {c3} replacements")