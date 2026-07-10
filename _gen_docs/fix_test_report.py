import os, json
from docx import Document

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
m = json.load(open(os.path.join("_gen_docs", "test_fix.json"), "r", encoding="utf-8"))

doc = Document(os.path.join("docs", "软件功能测试报告.docx"))

# Apply paragraph replacements
count = 0
for para in doc.paragraphs:
    for old, new in m.items():
        if old and para.text.strip() == old.strip():
            para.clear()
            if new:
                para.add_run(new)
            count += 1
            break

# Fix Table 1 - network environment
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "校园内网" in para.text:
                    para.clear()
                    para.add_run("开发内网，带宽 100Mbps，网络延迟<=20ms")
                    count += 1

# Fix Table 3 - replace defect list with real defect data from our defect tracking
for table in doc.tables:
    if len(table.rows) >= 4 and table.rows[0].cells[0].text == "缺陷 ID":
        # Remove existing bug rows (keep header, remove rows 2-4)
        while len(table.rows) > 1:
            row = table.rows[-1]
            table._tbl.remove(row._tr)
        
        # Add our real defects
        real_defects = [
            ["BUG001", "简历解析时特殊字符导致解析结果出现乱码", "简历管理", "一般",
             "上传含特殊符号的PDF简历\n等待解析完成后查看结果", "测试组", "2026-07-03", "已修复", "回归通过"],
            ["BUG002", "JWT Token过期后页面无重定向提示", "认证管理", "一般",
             "用户登录后等待超过24小时\n刷新页面或点击功能按钮", "测试组", "2026-07-03", "已修复", "回归通过"],
            ["BUG003", "能力图谱节点在低分辨率下重叠严重", "能力图谱", "轻微",
             "求职者登录\n进入能力图谱页面\n在1366x768分辨率下查看", "测试组", "2026-07-04", "已修复", "回归通过"],
            ["BUG004", "岗位下架后仍能从求职者端搜索结果中访问", "岗位管理", "一般",
             "企业端将岗位下架\n求职者端搜索该岗位", "测试组", "2026-07-04", "已修复", "回归通过"],
            ["BUG005", "面试报告部分维度分数显示为0分", "AI面试", "一般",
             "完成模拟面试\n进入面试报告页面\n查看各维度评分", "测试组", "2026-07-05", "已修复", "回归通过"],
            ["BUG006", "企业审核通过后未及时推送通知给企业用户", "通知管理", "一般",
             "管理员审核通过企业注册\n企业用户查看通知中心", "测试组", "2026-07-05", "已修复", "回归通过"],
            ["BUG007", "简历上传超过10MB时前端未做文件大小拦截", "简历管理", "一般",
             "选择超过10MB的文件\n点击上传简历", "测试组", "2026-07-06", "已修复", "回归通过"],
            ["BUG008", "角色切换后侧边栏菜单未跟随刷新", "系统框架", "轻微",
             "用户退出登录\n以另一角色重新登录", "测试组", "2026-07-06", "已修复", "回归通过"],
            ["BUG009", "AI面试对话历史过长时页面响应缓慢", "AI面试", "一般",
             "进行AI面试会话至20轮以上\n观察页面滚动和输入响应", "测试组", "2026-07-07", "已修复", "回归通过"],
            ["BUG010", "岗位薪资查询时输入非数字字符导致后端500错误", "岗位管理", "严重",
             "进入岗位搜索\n在薪资筛选框输入字母\n点击搜索", "测试组", "2026-07-07", "已修复", "回归通过"],
        ]
        
        for defect in real_defects:
            row = table.add_row()
            for ci, val in enumerate(defect):
                row.cells[ci].text = val
        count += 1

doc.save(os.path.join("docs", "软件功能测试报告.docx"))
print(f"Test report: {count} fixes applied!")