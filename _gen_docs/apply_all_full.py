# -*- coding: utf-8 -*-
import os, json
from docx import Document

os.chdir(r"C:\Users\Administrator\Desktop\Zhihire-StarMap")
base = os.path.join("docs")

def apply_maps(doc_path, *map_paths):
    doc = Document(os.path.join(base, doc_path))
    mapping = {}
    for mp in map_paths:
        mapping.update(json.load(open(os.path.join("_gen_docs", mp), "r", encoding="utf-8")))
    count = 0
    for para in doc.paragraphs:
        for old, new in mapping.items():
            if para.text.strip() == old.strip():
                para.clear()
                para.add_run(new)
                count += 1
                break
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in mapping.items():
                        if para.text.strip() == old.strip():
                            para.clear()
                            para.add_run(new)
                            count += 1
                            break
    doc.save(os.path.join(base, doc_path))
    return count

c1 = apply_maps("软件功能测试报告.docx", "test_map.json", "test_map_full.json")
print(f"Test report: {c1} replacements")

c2 = apply_maps("软件功能设计文档.docx", "design_map.json", "design_map_full.json")
print(f"Design doc: {c2} replacements")

c3 = apply_maps("软件产品说明书.docx", "prod_map.json")
print(f"Product manual: {c3} replacements")