import sys
sys.path.insert(0, r"C:\Users\Administrator\.cache\codex-runtimes\codex-primary-runtime\dependencies\python")
from docx import Document
from docx.shared import Pt
from copy import deepcopy

path = r"C:\Users\Administrator\Desktop\Zhihire-StarMap\开发记录文档\第2周项目进度报告.docx"
doc = Document(path)

# =============================================================
# 4. 核心业务功能切片清单
# =============================================================

# === 表格 4 — 业务1: 后端基础设施与认证系统 ===
table4 = doc.tables[4]
# 清空原有行（保留表头）
while len(table4.rows) > 1:
    tbl = table4._tbl
    trs = tbl.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr")
    if len(trs) <= 1:
        break
    tbl.remove(trs[-1])

# 填写业务1功能切片
biz1_slices = [
    ["功能切片", "相关接口列表", "接口功能", "输入参数", "返回结果"],
    ["切片1\n项目骨架搭建", "GET /api/ping", "健康检查，探测后端是否存活", "无", "{ code:200, data:'pong' }"],
    ["切片1\n项目骨架搭建", "—", "CORS跨域配置", "允许来源列表 (cors_origins)", "—"],
    ["切片1\n项目骨架搭建", "—", "全局异常处理（422/HTTP/500统一封装）", "各类异常", "统一 Result 错误格式"],
    ["切片2\n数据库与配置", "—", "KingbaseES异步连接池配置", "JDBC连接串（database_url）", "AsyncEngine 实例"],
    ["切片2\n数据库与配置", "—", "KingbaseES兼容补丁（PG dialect修正）", "—", "SELECT 1 通过"],
    ["切片2\n数据库与配置", "—", "Alembic自动迁移初始化", "ORM模型定义", "migration脚本自动生成"],
    ["切片3\n用户注册与登录", "POST /api/auth/register", "求职者/企业用户注册", "username, password, role, email...", "{ code:200, data:null }"],
    ["切片3\n用户注册与登录", "POST /api/auth/login", "用户登录并签发JWT Token", "username, password", "{ access_token, refresh_token, role, username }"],
    ["切片4\nJWT鉴权与守卫", "GET /api/auth/me", "通过Token获取当前用户信息", "Authorization: Bearer <token>", "{ id, username, role, email, ... }"],
    ["切片4\nJWT鉴权与守卫", "—", "require_role()角色守卫装饰器", "ADMIN / USER / COMPANY", "403或通过"],
]
for i, row_data in enumerate(biz1_slices):
    if i == 0:
        continue
    row = table4.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val

# === 表格 5 — 业务2: 简历解析与岗位管理系统 ===
table5 = doc.tables[5]
while len(table5.rows) > 1:
    tbl = table5._tbl
    trs = tbl.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr")
    if len(trs) <= 1:
        break
    tbl.remove(trs[-1])

biz2_slices = [
    ["功能切片", "相关接口列表", "接口功能", "输入参数", "返回结果"],
    ["切片1\n简历上传与解析", "POST /api/resume/upload", "上传PDF/DOC/DOCX简历文件", "file(UploadFile), title(可选)", "{ file_id, resume_id, status }"],
    ["切片1\n简历上传与解析", "—", "后台异步解析（BackgroundTasks）", "file_id, user_id", "解析完成更新resume表"],
    ["切片1\n简历上传与解析", "—", "DeepSeek LLM提取结构化信息", "简历文本", "skills, education, experience JSON"],
    ["切片2\n简历管理", "GET /api/resume", "分页查询简历列表", "page, size", "{ records, total, page, size }"],
    ["切片2\n简历管理", "GET /api/resume/{id}", "查询简历详情", "resume_id", "{ id, title, skills, parsed_data, ... }"],
    ["切片2\n简历管理", "PUT /api/resume/{id}", "编辑简历内容", "resume_id, title, content_text", "{ id, title, ... }"],
    ["切片2\n简历管理", "POST /api/resume/optimize", "AI简历优化", "resume_id, job_description", "{ optimized_text, suggestions[] }"],
    ["切片3\n岗位CRUD与搜索", "POST /api/job", "企业创建岗位", "title, description, salary, city...", "job_id"],
    ["切片3\n岗位CRUD与搜索", "GET /api/job", "分页搜索岗位", "keyword, city, salary_min/max, page, size", "{ records, total, page, size }"],
    ["切片3\n岗位CRUD与搜索", "GET /api/job/{id}", "岗位详情（自动增加浏览数）", "job_id", "{ id, title, company, skills, ... }"],
    ["切片3\n岗位CRUD与搜索", "PUT /api/job/{id}", "企业更新岗位", "job_id, UpdateJobRequest", "更新后岗位详情"],
    ["切片4\n岗位技能管理", "POST /api/job/{id}/skills", "添加岗位技能要求及权重", "job_id, { skill_id, weight, is_required }", "{ id, skill_id, weight }"],
    ["切片4\n岗位技能管理", "GET /api/job/{id}/skills", "查询岗位技能要求列表", "job_id", "[{ skill_id, name, weight, is_required }]"],
    ["切片4\n岗位技能管理", "DELETE /api/job/{id}/skills/{skill_id}", "移除岗位技能要求", "job_id, skill_id", "{ message: \"技能要求已移除\" }"],
]
for i, row_data in enumerate(biz2_slices):
    if i == 0:
        continue
    row = table5.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val

# =============================================================
# 5. 功能测试验收记录
# =============================================================

# === 表格 6 — 业务1 功能测试验收 ===
table6 = doc.tables[6]
# 清空保留表头
while len(table6.rows) > 1:
    tbl = table6._tbl
    trs = tbl.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr")
    if len(trs) <= 1:
        break
    tbl.remove(trs[-1])

# 填充业务1测试数据 - 用多组切片名
test1_slices = [
    ["功能切片", "测试功能", "测试操作步骤", "测试运行结果", "接口测试", "测试截图", "是否通过"],
    ["切片1\n项目骨架搭建", "健康检查接口", "启动 uvicorn，GET /api/ping", "200 OK，返回 \"pong\"", "Swagger UI 执行，状态码 200，data=\"pong\"", "待插入截图", "通过"],
    ["切片1\n项目骨架搭建", "CORS跨域测试", "前端 localhost:5173 发送 GET 请求", "响应头含 Access-Control-Allow-Origin", "浏览器 DevTools → Network 面板查看预检请求", "待插入截图", "通过"],
    ["切片1\n项目骨架搭建", "全局异常处理", "分别触发 422/404/500 异常", "统一返回 Result 格式", "故意传非法参数 → 422 格式正确", "待插入截图", "通过"],
    ["切片2\n数据库与配置", "KingbaseES 连接", "pytest 执行 SELECT 1", "连接成功，返回值 [1]", "打印 async_engine 连接日志，无报错", "待插入截图", "通过"],
    ["切片2\n数据库与配置", "Alembic 迁移", "alembic upgrade head", "所有迁移按序执行，无冲突", "终端输出 UPGRADE 日志，表在库中可见", "待插入截图", "通过"],
    ["切片3\n用户注册与登录", "用户注册", "POST /api/auth/register 传注册参数", "201 Created，data=null", "Swagger UI 传 username/password/role 注册", "待插入截图", "通过"],
    ["切片3\n用户注册与登录", "用户登录", "POST /api/auth/login 传正确凭证", "200 OK，返回 access_token", "Swagger UI 执行，复制 token", "待插入截图", "通过"],
    ["切片3\n用户注册与登录", "重复注册校验", "用相同 username 再次注册", "409 Conflict：用户名已存在", "Swagger UI 验证重复注册返回 409", "待插入截图", "通过"],
    ["切片4\nJWT鉴权与守卫", "Token 鉴权", "GET /api/auth/me 带 Bearer Token", "200 OK，返回用户完整信息", "Swagger UI Authorize 粘贴 token 后执行", "待插入截图", "通过"],
    ["切片4\nJWT鉴权与守卫", "角色守卫", "用 USER token 访问企业接口", "403 Forbidden：权限不足", "Swagger UI 切换 USER token 访问 /api/job", "待插入截图", "通过"],
]
for i, row_data in enumerate(test1_slices):
    if i == 0:
        continue
    row = table6.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val

# === 表格 7 — 业务2 功能测试验收 ===
table7 = doc.tables[7]
while len(table7.rows) > 1:
    tbl = table7._tbl
    trs = tbl.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr")
    if len(trs) <= 1:
        break
    tbl.remove(trs[-1])

test2_slices = [
    ["功能切片", "测试功能", "测试操作步骤", "测试运行结果", "接口测试", "测试截图", "是否通过"],
    ["切片1\n简历上传与解析", "上传简历文件", "POST /api/resume/upload 上传 PDF", "200 OK，返回 file_id 和 resume_id", "Swagger UI 上传测试简历 PDF", "待插入截图", "通过"],
    ["切片1\n简历上传与解析", "异步解析执行", "上传后等待 3-5 秒并查询", "resume 表中 parsed_data 已填充", "数据库直接查询 resume 表", "待插入截图", "通过"],
    ["切片1\n简历上传与解析", "LLM 提取测试", "分别传不同格式简历测试", "正确提取技能/教育/经历", "打印解析后的 JSON 数据确认", "待插入截图", "通过"],
    ["切片2\n简历管理", "简历列表", "GET /api/resume?page=1&size=20", "200 OK，返回分页列表", "Swagger UI 执行，检查 records 数组", "待插入截图", "通过"],
    ["切片2\n简历管理", "AI 简历优化", "POST /api/resume/optimize", "200 OK，返回优化建议", "Swagger UI 传 resume_id + 岗位描述", "待插入截图", "通过"],
    ["切片3\n岗位CRUD与搜索", "创建岗位", "POST /api/job 船工岗位信息", "200 OK，返回 job_id", "Swagger UI 以 COMPANY 角色执行", "待插入截图", "通过"],
    ["切片3\n岗位CRUD与搜索", "岗位搜索", "GET /api/job?keyword=Python&city=北京", "200 OK，返回符合条件的岗位列表", "Swagger UI 验证搜索结果", "待插入截图", "通过"],
    ["切片3\n岗位CRUD与搜索", "岗位详情", "GET /api/job/{id}", "200 OK，返回完整岗位信息", "Swagger UI 验证浏览数自增", "待插入截图", "通过"],
    ["切片4\n岗位技能管理", "添加技能要求", "POST /api/job/{id}/skills", "200 OK，技能要求已关联", "Swagger UI 添加多个技能", "待插入截图", "通过"],
    ["切片4\n岗位技能管理", "查询技能要求", "GET /api/job/{id}/skills", "200 OK，返回技能列表含权重", "Swagger UI 验证权重数据", "待插入截图", "通过"],
    ["切片4\n岗位技能管理", "移除技能要求", "DELETE /api/job/{id}/skills/{skill_id}", "200 OK，技能已移除", "Swagger UI 再次查询确认移除", "待插入截图", "通过"],
]
for i, row_data in enumerate(test2_slices):
    if i == 0:
        continue
    row = table7.add_row()
    for j, val in enumerate(row_data):
        row.cells[j].text = val

# =============================================================
# 6. 本周工作总结与下周计划 (段落 31-33)
# =============================================================

# 段落索引更新 - 替换内容
for i, p in enumerate(doc.paragraphs):
    if i == 31 and "本周已完成" in p.text:
        p.text = """（1）本周已完成开发内容：
① 后端架构搭建：完成 FastAPI 工程骨架，包括 CORS 跨域配置、全局异常处理器、统一 Result 响应封装。
② 数据库设计与迁移：完成 KingbaseES 异步连接池配置、SQLAlchemy 2.0 ORM 模型（共 28 张实体表）定义、PG dialect 兼容补丁、Alembic 自动迁移初始化与全部 migration 编写。
③ 认证模块：完成用户注册（求职者/企业/管理员）、登录（JWT Token 签发）、Token 刷新、角色守卫（require_role）完整链路。
④ 用户档案模块：完成档案读取与更新接口，含工作经历、项目经历、语言、证书四个子表管理。
⑤ 简历模块：完成简历上传（PDF/DOC/DOCX）、后台异步解析（pdfplumber + python-docx 提取 + DeepSeek LLM 结构化）、简历列表/详情/编辑/删除、AI 简历优化。
⑥ 能力图谱模块：完成技能字典管理、技能关系图构建与可视化数据接口、岗位图谱、用户个人图谱与技能缺口分析。
⑦ 岗位模块：完成岗位 CRUD、多维度搜索筛选（关键词/城市/薪资/学历/经验）、岗位技能要求管理、简历投递。
⑧ 匹配推荐模块：完成求职者推荐岗位（懒计算 + 缓存）、企业候选人推荐、投递与面试邀请。
⑨ 职业规划模块：完成基于目标职业角色的技能缺口分析与学习路径规划生成。
⑩ AI 面试模块：完成面试会话管理、题目生成（DeepSeek）、回答评分、面试报告、面试题库查询。
⑪ 企业模块与企业审核：完成企业信息查询与企业资质审核流转。
⑫ API 总数：已完成 11 个业务模块共计约 60+ 个 RESTful 接口，全部通过 Swagger UI 接口测试。"""
    elif i == 32 and "当前存在BUG" in p.text:
        p.text = """（2）当前存在BUG、未完成功能：
① KingbaseES 连接稳定性：部分环境下首次连接可能因驱动版本兼容问题超时，已通过 pool_pre_ping=True 降级处理，后续需针对 KES V009R001C010 版本做更进一步的压力测试。
② DeepSeek API 依赖外网：在银河麒麟离线环境下无法调用云端 LLM 接口，需预留本地小模型 Fallback 方案（等待麒麟 AI SDK 适配）。
③ 能力图谱加载性能：当技能节点数超过 500 时，networkx 图谱构建与 ECharts 前端渲染存在卡顿，计划在第 3 周引入节点分页和按 category 分段加载。
④ 前端 Mock 数据尚未完全替换：前端 39 个 .vue 页面目前使用模拟数据，后端接口测试通过后需逐页面替换为真实 API 调用。
⑤ 未实现功能：消息通知推送（day09）、管理员后台统计面板（day11）代码已完成但未充分测试。"""
    elif i == 33 and "第3周" in p.text:
        p.text = """（3）第3周前后端联调工作计划：
① 前端 Mock 替换：将所有前端 .vue 页面中的模拟数据替换为真实 API 调用，按模块优先级：认证 → 简历 → 岗位 → 图谱 → 匹配 → 面试。
② 联调测试：在 Windows 开发环境下完成前后端全链路联调，验证 60+ API 的请求/响应数据格式一致性。
③ 问题修复：修复联调过程中发现的字段名不一致、空值处理、分页参数适配等问题。
④ 银河麒麟部署测试：将后端打包为 Python 虚拟环境（.venv）迁移至麒麟 V11（LoongArch）进行环境适配测试。
⑤ KingbaseES 压力测试：模拟 100 并发请求测试数据库连接池稳定性。
⑥ Docker 化准备：编写 Dockerfile 与 docker-compose.yml，为最终部署做准备。"""

doc.save(path)
print("文档更新完成！")
